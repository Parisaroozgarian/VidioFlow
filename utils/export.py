import io
import re
import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_script(script_text):
    section_pattern = r"(#{1,3}|[A-Z\s]+:)\s*(.*)"
    visual_pattern = r"\[VISUAL[^\]]*\](.*?)(?=\[|$)"
    caption_pattern = r"\[CAPTION[^\]]*\](.*?)(?=\[|$)"

    lines = script_text.split('\n')
    sections = []
    current_section = {"title": "Script", "content": [], "visuals": [], "captions": []}

    for line in lines:
        section_match = re.match(section_pattern, line)
        if section_match:
            if current_section["content"]:
                sections.append(current_section)
                current_section = {"title": section_match.group(2), "content": [], "visuals": [], "captions": []}
            else:
                current_section["title"] = section_match.group(2)
            continue

        visual_match = re.search(visual_pattern, line)
        if visual_match:
            current_section["visuals"].append(visual_match.group(1).strip())
            continue

        caption_match = re.search(caption_pattern, line)
        if caption_match:
            current_section["captions"].append(caption_match.group(1).strip())
            continue

        if line.strip():
            current_section["content"].append(line)

    if current_section["content"]:
        sections.append(current_section)

    return sections


def horizontal_line():
    return Table([['']], colWidths=[450],
                 style=[('LINEABOVE', (0, 0), (-1, -1), 1, colors.black)])


def generate_pdf(script_text, title="Video Script"):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='SectionHeader',
                              parent=styles['Heading2'],
                              spaceAfter=6))
    styles.add(ParagraphStyle(name='NormalIndent',
                              parent=styles['Normal'],
                              leftIndent=20,
                              spaceAfter=6))

    elements = []
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    today = datetime.datetime.now().strftime("%B %d, %Y")
    elements.append(Paragraph(f"Generated on {today}", styles['Normal']))
    elements.append(Spacer(1, 24))

    for section in parse_script(script_text):
        elements.append(Paragraph(section["title"], styles['SectionHeader']))
        for line in section["content"]:
            elements.append(Paragraph(line, styles['Normal']))

        if section["visuals"]:
            elements.append(Spacer(1, 6))
            elements.append(Paragraph("Visual Notes:", styles['Italic']))
            for visual in section["visuals"]:
                elements.append(Paragraph(f"• {visual}", styles['NormalIndent']))

        if section["captions"]:
            elements.append(Spacer(1, 6))
            elements.append(Paragraph("Caption Suggestions:", styles['Italic']))
            for caption in section["captions"]:
                elements.append(Paragraph(f"• {caption}", styles['NormalIndent']))

        elements.append(Spacer(1, 12))
        elements.append(horizontal_line())
        elements.append(Spacer(1, 12))

    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_docx(script_text, title="Video Script"):
    document = Document()
    buffer = io.BytesIO()

    document.add_heading(title, level=0)

    today = datetime.datetime.now().strftime("%B %d, %Y")
    date_paragraph = document.add_paragraph(f"Generated on {today}")
    date_paragraph.style = 'Subtitle'

    document.add_paragraph()

    for section in parse_script(script_text):
        document.add_heading(section["title"], level=1)
        document.add_paragraph("\n".join(section["content"]))

        if section["visuals"]:
            document.add_paragraph("Visual Notes:", style='Intense Quote')
            for visual in section["visuals"]:
                document.add_paragraph(visual, style='List Bullet')

        if section["captions"]:
            document.add_paragraph("Caption Suggestions:", style='Intense Quote')
            for caption in section["captions"]:
                document.add_paragraph(caption, style='List Bullet')

        document.add_paragraph()

    document.save(buffer)
    buffer.seek(0)
    return buffer
# Export util
